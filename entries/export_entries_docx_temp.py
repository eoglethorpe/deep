import docx
from docx import RT
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from entries.models import *


# See: https://github.com/python-openxml/python-docx/issues/74#issuecomment-215678765
def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    r = docx.text.run.Run(new_run, paragraph)

    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


# From https://github.com/python-openxml/python-docx/issues/105
def add_line(para):
    # a very convoluded way to add a horizontal line to the document

    def first_child_found_in(parent, tagnames):
        """
        Return the first child of parent with tag in *tagnames*, or None if
        not found.
        """
        for tagname in tagnames:
            child = parent.find(qn(tagname))
            if child is not None:
                return child
        return None

    def insert_element_before(parent, elm, successors):
        """
        Insert *elm* as child of *parent* before any existing child having
        tag name found in *successors*.
        """
        successor = first_child_found_in(parent, successors)
        if successor is not None:
            successor.addprevious(elm)
        else:
            parent.append(elm)
        return elm

    p = para._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    insert_element_before(pPr, pBdr, successors=(
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    ))
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)


def add_excerpt_info(d, info):
    # Show the excerpt
    ref = d.add_paragraph(info.excerpt)
    ref.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    # Show the reference

    source_name = ""
    if info.entry.lead.source and info.entry.lead.source.name != "":
        source_name = info.entry.lead.source.name
    
    if source_name == "":
        source_name = "Reference"

    ref.add_run(" (")
    if info.entry.lead.url and info.entry.lead.url != "":
        add_hyperlink(ref, info.entry.lead.url, source_name)
    
    elif Attachment.objects.filter(lead=info.entry.lead).count() > 0:
        add_hyperlink(ref, info.entry.lead.attachment.upload.url, source_name)

    else:
        ref.add_run("Manual Entry")

    # TODO Find out whether to show date of info or lead
    if info.date:
        ref.add_run(", {}".format(info.date.strftime("%m/%d/%Y")), )


    # d.add_paragraph("Reliability: {}\nSeverity: {}".format(info.reliability.name, info.severity.name))

    r = ref.add_run()
    r.add_text(' ')
    r.add_picture('static/img/doc_export/sev_{}.png'.format(info.severity.level), height=docx.shared.Inches(.17))
    r.add_text(' ')
    r.add_picture('static/img/doc_export/rel_{}.png'.format(info.reliability.level), height=docx.shared.Inches(.17))
    r.add_text(' ')

    ref.add_run(")")

    d.add_paragraph()


def set_style(style):
    # style.paragraph_format.space_after = docx.shared.Pt(6)
    style.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT


def export_docx(order, event):
    d = docx.Document('static/doc_export/template.docx')

    # Set document styles
    set_style(d.styles["Normal"])
    set_style(d.styles["Heading 1"])
    set_style(d.styles["Heading 2"])
    set_style(d.styles["Heading 3"])
    set_style(d.styles["Heading 4"])
    set_style(d.styles["Heading 5"])
    

    # TODO: Hierarchy and filter

    # The leads for which excerpts we displayed
    leads_pk = []

    # First the attributes with no sectors

    # Get each pillar
    pillars = InformationPillar.objects.filter(contains_sectors=False)
    for pillar in pillars:
        pillar_header_shown = False
        
        # Get each subpillar
        subpillars = pillar.informationsubpillar_set.all()
        for subpillar in subpillars:
            attributes = InformationAttribute.objects.filter(subpillar=subpillar,
                                                             sector=None,
                                                             information__entry__lead__event__pk=event)

            if len(attributes) > 0:
                if not pillar_header_shown:
                    d.add_heading(pillar.name, level=2)
                    d.add_paragraph()
                    pillar_header_shown = True
                d.add_heading(subpillar.name, level=3)
                d.add_paragraph()

            for attr in attributes:
                info = attr.information
                add_excerpt_info(d, info)
                leads_pk.append(info.entry.lead.pk)

    # Next the attributes containing sectors

    # Get each sector
    for sector in Sector.objects.all():
        sector_header_shown = False

        # Get each pillar
        pillars = InformationPillar.objects.filter(contains_sectors=True)
        for pillar in pillars:
            pillar_header_shown = False

            # Get each subpillar
            subpillars = pillar.informationsubpillar_set.all()
            for subpillar in subpillars:
                attributes = InformationAttribute.objects.filter(subpillar=subpillar,
                                                                 sector=sector,
                                                                 information__entry__lead__event__pk=event)

                if len(attributes) > 0:
                    if not sector_header_shown:
                        d.add_heading(sector.name, level=2)
                        d.add_paragraph()
                        sector_header_shown = True
                    if not pillar_header_shown:
                        d.add_heading(pillar.name, level=3)
                        d.add_paragraph()
                        pillar_header_shown = True
                    d.add_heading(subpillar.name+":", level=4)

                for attr in attributes:
                    info = attr.information
                    add_excerpt_info(d, info)
                    leads_pk.append(info.entry.lead.pk)

    
    add_line(d.add_paragraph())

    # Bibliography
    d.add_paragraph()
    h1 = d.add_heading("Bibliography", level=1)
    d.add_paragraph()

    leads_pk = list(set(leads_pk))
    leads = Lead.objects.filter(pk__in=leads_pk)
    for lead in leads:
        p = d.add_paragraph()
        if lead.source:
            p.add_run(lead.source.name.title())
        else:
            p.add_run("Missing source".title())

        p.add_run(". {}.".format(lead.name.title()))
        if lead.published_at:
            p.add_run(" {}.".format(lead.published_at.strftime("%m/%d/%Y")))

        p = d.add_paragraph()
        if lead.url and lead.url != "":
            add_hyperlink(p, lead.url, lead.url)
        
        elif Attachment.objects.filter(lead=lead).count() > 0:
            add_hyperlink(p, lead.attachment.upload.url, lead.attachment.upload.url)

        else:
            p.add_run("Missing url.")
            
        d.add_paragraph()

    d.add_page_break()

    return d