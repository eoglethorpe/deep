import string
import requests
import datetime
import tempfile
import base64
import re

import docx
# from docx import RT
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
# from docx.enum.style import WD_STYLE_TYPE

from entries import models as entry_model
from leads import models as lead_model
from entries.strippers import write_file


def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF or
        codepoint in (0x9, 0xA, 0xD) or
        0xE000 <= codepoint <= 0xFFFD or
        0x10000 <= codepoint <= 0x10FFFF
    )


def xstr(conv):
    """remove illegal characters from a string (errors from PDFs etc)"""
    s = "".join(filter(lambda x: x in string.printable, conv))
    return ''.join(c for c in s if valid_xml_char_ordinal(c))


# See:
# https://github.com/python-openxml/python-docx/issues/74#issuecomment-215678765
def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file
    # and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
                          is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add
    # add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    r = docx.text.run.Run(new_run, paragraph)

    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


# From https://github.com/python-openxml/python-docx/issues/105
def add_line(para):
    # a very convoluded way to add a horizontal line to the document

    def first_child_found_in(parent, tagnames):
        """
        Return the first child of parent with tag in *tagnames*, or None if
        not found.
        """
        for tagname in tagnames:
            child = parent.find(qn(tagname))
            if child is not None:
                return child
        return None

    def insert_element_before(parent, elm, successors):
        """
        Insert *elm* as child of *parent* before any existing child having
        tag name found in *successors*.
        """
        successor = first_child_found_in(parent, successors)
        if successor is not None:
            successor.addprevious(elm)
        else:
            parent.append(elm)
        return elm

    p = para._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    insert_element_before(pPr, pBdr, successors=(
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    ))
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)


def add_excerpt_info(d, info):
    # Show the excerpt
    try:
        ref = d.add_paragraph(xstr(info.excerpt))
        if len(info.image):
            fimage = tempfile.NamedTemporaryFile()
            if re.search(r'http[s]?://', info.image):
                image = requests.get(info.image, stream=True)
                write_file(image, fimage)
            else:
                image = base64.b64decode(info.image.split(',')[1])
                fimage.write(image)
            d.add_picture(fimage)
    except:
        ref = d.add_paragraph('')
    ref.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    # Show the reference

    source_name = ""
    if info.entry.lead.source_name and info.entry.lead.source_name != "":
        source_name = info.entry.lead.source_name

    if source_name == "":
        source_name = "Reference"

    ref.add_run(" (")
    if info.entry.lead.url and info.entry.lead.url != "":
        add_hyperlink(ref, info.entry.lead.url, source_name)

    elif entry_model.Attachment.objects.filter(
            lead=info.entry.lead).count() > 0:
        add_hyperlink(ref, info.entry.lead.attachment.upload.url, source_name)

    else:
        ref.add_run("Manual Entry")

    # TODO Find out whether to show date of info or lead
    if info.date:
        ref.add_run(", {}".format(info.date.strftime("%d/%m/%Y")), )

    # d.add_paragraph("Reliability: {}\nSeverity: {}"
    #                  .format(info.reliability.name, info.severity.name))

    r = ref.add_run()
    r.add_text(' ')
    r.add_picture('static/img/doc_export/sev_{}.png'
                  .format(info.severity.level), height=docx.shared.Inches(.17))
    r.add_text(' ')
    r.add_picture('static/img/doc_export/rel_{}.png'
                  .format(info.reliability.level),
                  height=docx.shared.Inches(.17))
    r.add_text(' ')

    ref.add_run(")")

    d.add_paragraph()


def set_style(style):
    # style.paragraph_format.space_after = docx.shared.Pt(6)
    style.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT


def export_docx(event, informations=None, export_geo=False):
    d = docx.Document('static/doc_export/template.docx')

    # Set document styles
    set_style(d.styles["Normal"])
    set_style(d.styles["Heading 1"])
    set_style(d.styles["Heading 2"])
    set_style(d.styles["Heading 3"])
    set_style(d.styles["Heading 4"])
    set_style(d.styles["Heading 5"])

    # The leads for which excerpts we displayed
    leads_pk = []

    # First the attributes with no sectors

    # Get each pillar
    pillars = entry_model.InformationPillar.objects.filter(
                contains_sectors=False)
    for pillar in pillars:
        pillar_header_shown = False

        # Get each subpillar
        subpillars = pillar.informationsubpillar_set.all()
        for subpillar in subpillars:
            attributes = entry_model.InformationAttribute.objects.filter(
                    subpillar=subpillar,
                    sector=None,
                    information__entry__lead__event__pk=event)
            if informations is not None:
                attributes = attributes.filter(
                        information__pk__in=informations)

            if len(attributes) > 0:
                if not pillar_header_shown:
                    d.add_heading(pillar.name, level=2)
                    d.add_paragraph()
                    pillar_header_shown = True
                d.add_heading(subpillar.name, level=3)
                d.add_paragraph()

            already_added = []
            for attr in attributes:
                info = attr.information
                if info not in already_added:
                    already_added.append(info)
                    add_excerpt_info(d, info)
                    leads_pk.append(info.entry.lead.pk)

    # Next the attributes containing sectors

    # Get each sector
    for sector in entry_model.Sector.objects.all():
        sector_header_shown = False

        if export_geo:
            attributes = entry_model.InformationAttribute.objects.filter(
                    sector=sector,
                    information__entry__lead__event__pk=event)
            if informations is not None:
                attributes = attributes.filter(
                        information__pk__in=informations)

            if len(attributes) > 0:
                if not sector_header_shown:
                    d.add_heading(sector.name, level=2)
                    d.add_paragraph()
                    sector_header_shown = True

            already_added = []
            for attr in attributes:
                info = attr.information
                if info not in already_added:
                    already_added.append(info)
                    add_excerpt_info(d, info)
                    leads_pk.append(info.entry.lead.pk)
        else:
            # Get each pillar
            pillars = entry_model.InformationPillar.objects.filter(
                        contains_sectors=True)
            for pillar in pillars:
                pillar_header_shown = False

                # Get each subpillar
                subpillars = pillar.informationsubpillar_set.all()
                for subpillar in subpillars:
                    attributes = entry_model.InformationAttribute.objects.filter(
                            subpillar=subpillar,
                            sector=sector,
                            information__entry__lead__event__pk=event)
                    if informations is not None:
                        attributes = attributes.filter(
                                information__pk__in=informations)

                    if len(attributes) > 0:
                        if not sector_header_shown:
                            d.add_heading(sector.name, level=2)
                            d.add_paragraph()
                            sector_header_shown = True
                        if not pillar_header_shown:
                            d.add_heading(pillar.name, level=3)
                            d.add_paragraph()
                            pillar_header_shown = True
                        d.add_heading(subpillar.name+":", level=4)

                    already_added = []
                    for attr in attributes:
                        info = attr.information
                        if info not in already_added:
                            already_added.append(info)
                            add_excerpt_info(d, info)
                            leads_pk.append(info.entry.lead.pk)

    add_line(d.add_paragraph())

    # Bibliography
    d.add_paragraph()
    d.add_heading("Bibliography", level=1)
    d.add_paragraph()

    if len(leads_pk) == 0:
        leads = []
    else:
        leads_pk = list(set(leads_pk))
        leads = entry_model.Lead.objects.filter(pk__in=leads_pk)

    for lead in leads:
        p = d.add_paragraph()
        if lead.source_name and lead.source_name != "":
            p.add_run(lead.source_name.title())
        else:
            p.add_run("Missing source".title())

        p.add_run(". {}.".format(lead.name.title()))
        if lead.published_at:
            p.add_run(" {}.".format(lead.published_at.strftime("%m/%d/%Y")))

        p = d.add_paragraph()
        if lead.url and lead.url != "":
            add_hyperlink(p, lead.url, lead.url)

        elif entry_model.Attachment.objects.filter(lead=lead).count() > 0:
            add_hyperlink(p, lead.attachment.upload.url,
                          lead.attachment.upload.url)

        else:
            p.add_run("Missing url.")

        d.add_paragraph()

    d.add_page_break()

    return d


def export_docx_new_format(event, informations=None):
    """
    Export As Specified in Issue

    #259
    New Export format in word - Briefing note template
    """

    d = docx.Document('static/doc_export/template_02.docx')

    def _filter_info(xyz_info):
        """
        Filter with given `informations` if provided
        """
        if informations is not None:
            xyz_info = xyz_info.filter(
                    information__pk__in=informations)
        return xyz_info

    def add_info(xyz_info):
        xyz_info = _filter_info(xyz_info)
        for info in xyz_info.all():
            add_excerpt_info(d, info.information)

    def add_heading(heading, level):
        d.add_paragraph()
        d.add_heading(heading, level=level)
        d.add_paragraph()

    section = d.sections[0]

    # Already Changed in the template(template_02.docx)
    """
    # Page Orientation Change
    section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
    new_width = section.page_width
    section.page_width = section.page_height
    section.page_height = new_width

    # Divide Page into 2 col
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), "2")
    """

    # Set document styles
    set_style(d.styles["Normal"])
    set_style(d.styles["Heading 1"])
    set_style(d.styles["Heading 2"])
    set_style(d.styles["Heading 3"])
    set_style(d.styles["Heading 4"])
    set_style(d.styles["Heading 5"])

    # For Blank paragraph
    BLANK = '--\n\n\n\n--'

    # Define Style (Already Definded in Template)
    """
    d.styles.add_style('Title Briefing Text',
                       WD_STYLE_TYPE.PARAGRAPH)
    d.styles.add_style('Title Briefing Time',
                       WD_STYLE_TYPE.CHARACTER)
    d.styles.add_style('Title Country/Crisis Name',
                       WD_STYLE_TYPE.PARAGRAPH)
    d.styles.add_style('Notice',
                       WD_STYLE_TYPE.PARAGRAPH)
    d.styles.add_style('Limitations',
                       WD_STYLE_TYPE.PARAGRAPH)
    d.styles.add_style('Sector Impact Sector Name',
                       WD_STYLE_TYPE.PARAGRAPH)
    """

    #                   Event/Crisis object
    event = lead_model.Event.objects.get(pk=event)

    p = d.add_paragraph('Briefing Note – ', style='Title Briefing Text')
    p.add_run(datetime.datetime.now().strftime("%m %b %Y"),
              style='Title Briefing Time')

    # Already Added in the Template
    """
    d.add_paragraph().add_run(style='Title Logo')\
        .add_picture('static/doc_export/acaps_logo.png',
                     docx.shared.Inches(.75),
                     docx.shared.Inches(.56))
    d.paragraphs[-1].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
    d.add_paragraph().add_run(style='Title Logo')\
        .add_picture('static/doc_export/startNetwork_logo.png',
                     docx.shared.Inches(1.05),
                     docx.shared.Inches(.51))
    d.paragraphs[-1].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
    """

    d.add_paragraph(
            ', '.join([country[0]
                       for country in event.countries.values_list('name')]),
            style='Title Country/Crisis Name')
    d.add_paragraph(event.name, style='Title Country/Crisis Name')

    d.add_paragraph()

    # Table after country and crisis Name
    rows = [['Need for international assistance', 'Not required',
             'Low', 'Moderate', 'Significant', 'Urgent'],
            ['', '#FEFFB1', '#FECC5D', '#FD8C3E', '#F03B20', '#BD0026'],
            ['Expected impact', 'Insignificant', 'Minor', 'Moderate',
             'Significant', 'Major'],
            ['', '#FEFFB1', '#FECC5D', '#FD8C3E', '#F03B20', '#BD0026']]
    table = d.add_table(rows=len(rows), cols=len(rows[0]))
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    import re
    for index_row, row in enumerate(rows):
        for index_col, cell_value in enumerate(row):
            cell = table.cell(index_row, index_col)
            if re.match('^#[\w\d]+$', cell_value):
                shading_elm = parse_xml(
                        r'<w:shd {} w:fill="{}"/>'
                        .format(nsdecls('w'), cell_value.replace('#', '')))
                cell._tc.get_or_add_tcPr().append(shading_elm)
            else:
                cell.text = str(cell_value)
        if index_row % 2 == 1:
            table.cell(index_row, 0).merge(table.cell(index_row-1, 0))
    d.add_paragraph()

    #                   Crisis overview
    add_heading('Crisis overview', level=2)
    context_info = entry_model.InformationAttribute.objects.filter(
            subpillar__name__iregex=r'(Overview|developments)',
            subpillar__pillar__name__iexact='Context',
            information__entry__lead__event=event
            ).all()

    add_info(context_info)

    cols = ['Affected groups', 'Affected area', 'Total figures',
            '% of pop. affected']
    rows = ['Resident pop.', 'Killed  ', 'Injured',
            'Missing', 'Total IDPs', 'Total refugees',
            'Total displaced', 'Etc.']
    table = d.add_table(rows=len(rows)+4, cols=len(cols))
    table.style = 'Light Shading'
    for index, col in enumerate(cols):
        table.cell(0, index).text = str(col)
    for index, row in enumerate(rows):
        table.cell(index+1, 0).text = str(row)
        if row == 'Etc.':
            table.cell(index+2, 0).merge(table.cell(index+1, 0))

    #                   Key findings
    d.add_paragraph()
    add_heading('Key findings', level=2)

    # Table
    d.add_paragraph()
    rows = ['Anticipated scope and scale',
            'Priorities for humanitarian intervention',
            'Humanitarian constraints']
    table = d.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = rows[0]
    for row in rows[1:]:
        row_cells = table.add_row().cells
        row_cells[0].text = str(row)

    d.add_paragraph()
    p = d.add_paragraph(style='Limitations')
    charac = p.add_run('Limitations')
    charac.italic = True
    charac.bold = True
    p.add_run('\nDetail any limitations due to '
              'missing/suspect data, issues with the scope of the'
              ' report or with the findings.')
    d.add_paragraph()

    #                   Crisis impact
    add_heading('Crisis impact', level=2)
    d.add_heading('Overview', level=3)
    context_info = entry_model.InformationAttribute.objects.filter(
            subpillar__pillar__name__iregex=r'(Humanitarian conditions)',
            sector__name__iexact='Cross',
            information__entry__lead__event=event
            )
    add_info(context_info)

    d.add_heading('Humanitarian profile', level=3)
    context_info = entry_model.InformationAttribute.objects.filter(
            subpillar__name__iregex=r'(Humanitarian profile|'
                                    'population displacement)',
            subpillar__pillar__name__iexact='Population Profile',
            information__entry__lead__event=event
            )
    add_info(context_info)

    d.add_heading('Sectoral impact', level=3)
    sectors = ['Food', 'Livelihood', 'WASH', 'Health', 'Shelter',
               'NFI', 'Protection', 'Education', 'Nutrition',
               'Agriculture']
    for sector in sectors:
        d.add_paragraph(sector+' :', style='Sector Impact Sector Name')
        sector_info = entry_model.InformationAttribute.objects.filter(
            subpillar__pillar__name__iregex=r'(Humanitarian conditions)',
            sector__name__iexact=sector,
            information__entry__lead__event=event
            )
        add_info(sector_info)

    d.add_heading('Impact on critical infrastructure', level=3)
    context_info = entry_model.InformationAttribute.objects.filter(
            subpillar__name__iregex=r'(Systems disruption|Damages|Losses)',
            subpillar__pillar__name__iexact='Scope and scale',
            information__entry__lead__event=event
            )
    add_info(context_info)

    d.add_heading('Vulnerable groups affected', level=3)
    context_info = entry_model.InformationAttribute.objects.filter(
            subpillar__name__iregex=r'(vulnerabilities|risks)',
            subpillar__pillar__name__iexact='Humanitarian conditions',
            information__entry__lead__event=event
            )
    add_info(context_info)

    context_info = entry_model.InformationAttribute.objects.filter(
            subpillar__name__iexact='Population with specific needs',
            subpillar__pillar__name__iexact='Population profile',
            information__entry__lead__event=event
            )
    add_info(context_info)

    d.add_heading('Humanitarian and operational constraints', level=3)
    subpillars = ['Humanitarian access gaps', 'Relief to beneficiaries',
                  'Beneficiaries to relief', 'Physical Constraints']
    for subpillar in subpillars:
        d.add_paragraph(subpillar+' :', style='Sector Impact Sector Name')
        context_info = entry_model.InformationAttribute.objects.filter(
                subpillar__name__iexact=subpillar,
                subpillar__pillar__name__iexact='Humanitarian access',
                information__entry__lead__event=event
                )
        add_info(context_info)

    #                   Potential aggravating factors
    add_heading('Potential aggravating factors', level=2)

    d.add_heading('Seasonal information', level=3)
    d.add_paragraph(BLANK)

    d.add_heading('Other factors of vulnerability', level=3)
    context_info = entry_model.InformationAttribute.objects.filter(
            subpillar__pillar__name__iregex=r'(Humanitarian conditions|risks)',
            sector__name__iexact='Cross',
            information__entry__lead__event=event
            )
    add_info(context_info)

    context_info = entry_model.InformationAttribute.objects.filter(
            subpillar__name__iregex=r'(drivers|aggravating factors)',
            subpillar__pillar__name__iexact='Context',
            information__entry__lead__event=event
            )
    add_info(context_info)

    #                   Contextual information
    add_heading('Contextual information', level=2)

    d.add_heading('Drivers of the current conflict', level=3)
    context_info = entry_model.InformationAttribute.objects.filter(
            subpillar__name__iregex=r'(Society and community|politics '
                                    'and security|hazard developments)',
            subpillar__pillar__name__iexact='Context',
            information__entry__lead__event=event)
    add_info(context_info)

    d.add_heading('Relevant stakeholders', level=3)
    context_info = entry_model.InformationAttribute.objects.filter(
            subpillar__name__iexact='Stakeholders',
            subpillar__pillar__name__iexact='Context',
            information__entry__lead__event=event)
    add_info(context_info)

    d.add_heading('International and neighbouring countries’ relationship '
                  'to the conflict ', level=3)
    d.add_paragraph(BLANK)

    d.add_heading('Past disasters or displacement', level=3)
    d.add_paragraph('Graph might appear here!!')

    #                   Key characteristics
    add_heading('Key characteristics ', level=2)
    cols = ['Demographic profile ', 'WASH statistics', 'Lighting and cooking',
            'Key health statistics', 'Food', 'Nutrition', 'Literacy']
    rows = ['Total population in country affected areas, gender'
            ' and age distribution and rural vs urban.',
            'Statistics: (access to improved sources of drinking water'
            ', access to toilet facilities).',
            'sources.', 'Infant mortality rate, under-5 mortality rate,'
            ' maternal ', 'Mortality rate.', 'Figures.', 'Levels.', 'Levels.'
            ]
    for index, col in enumerate(cols):
        p = d.add_paragraph(style="List Bullet")
        p.add_run(col).bold = True
        p.add_run(' '+rows[index]+'\n')

    d.add_paragraph(
            "Sources: National statistical systems, Global Population "
            "Statistics, UNFPA country profiles, Population Stats, CIA"
            " World Factbook, UNICEF/WHO JMP, WFP food security reports,"
            " UNICEF country statistics, UNWATER country overview, GIEWS"
            " country briefs, UN Standing Committee on Nutrition, WHO"
            " Nutrition country profiles, WHO country statistics, UNICEF"
            " state of world’s children, UNFPA country profiles,"
            " World Bank Databank.")

    #                   Response capacity
    add_heading('Response capacity', level=2)

    context_info = entry_model.InformationAttribute.objects.filter(
            subpillar__name__iexact='response gaps',
            subpillar__pillar__name__iexact='Capacities and response',
            information__entry__lead__event=event)
    add_info(context_info)

    d.add_heading('Local and national response capacity', level=3)
    context_info = entry_model.InformationAttribute.objects.filter(
            subpillar__name__iexact='National response',
            subpillar__pillar__name__iexact='Capacities and response',
            information__entry__lead__event=event)
    add_info(context_info)

    d.add_heading('International response capacity', level=3)
    context_info = entry_model.InformationAttribute.objects.filter(
            subpillar__name__iexact='International response',
            subpillar__pillar__name__iexact='Capacities and response',
            information__entry__lead__event=event)
    add_info(context_info)

    d.add_heading('Population coping mechanisms', level=3)
    context_info = entry_model.InformationAttribute.objects.filter(
            subpillar__name__iexact='Copying mechanisms',
            subpillar__pillar__name__iexact='Capacities and response',
            information__entry__lead__event=event)
    add_info(context_info)

    d.add_heading('Market functionality', level=3)
    context_info = entry_model.InformationAttribute.objects.filter(
            subpillar__name__iexact='Market functionality',
            subpillar__pillar__name__iexact='Capacities and response',
            information__entry__lead__event=event)
    add_info(context_info)

    #                   Information gaps and needs
    add_heading('Information gaps and needs', level=2)

    context_info = entry_model.InformationAttribute.objects.filter(
            subpillar__name__iregex='(Communication means|information '
                                    'challenges|information needs'
                                    '|Information gaps)',
            subpillar__pillar__name__iexact='Communication',
            information__entry__lead__event=event)
    add_info(context_info)

    #                   Lessons learned
    add_heading('Lessons learned', level=2)

    subpillars = ['Humanitarian conditions', 'Scope and Scale',
                  'Capacities and response', 'Context']

    for subpillar in subpillars:
        d.add_heading(subpillar, level=3)
        context_info = entry_model.InformationAttribute.objects.filter(
                subpillar__name__iexact='lessons learnt',
                subpillar__pillar__name__iexact=subpillar,
                information__entry__lead__event=event)
        add_info(context_info)

    d.add_page_break()

    #   #   #   #   #   #   #   #
    # Divide Page into 1 col
    section = d.add_section()
    sectPr = section._sectPr
    # Single Col
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), "1")
    # Width Fix
    col_width = sectPr.xpath('./w:cols/w:col')[0]
    col_width.set(qn('w:w'), "12534")
    col_width.set(qn('w:space'), "720")

    #                   Key characteristics
    add_heading('Key characteristics', level=2)

    cols = ['Key indicators', 'Area 1', 'Area 2', 'Area 3']
    rows = ['Total population', '% population in rural areas',
            'Gender and age distribution of population',
            'State capital', 'Lighting and cooking sources',
            'WASH figures', 'Health figures',
            'Food security', 'Nutrition levels',
            'Literacy rates', 'Others']
    d.add_paragraph()
    table = d.add_table(rows=1, cols=len(cols))
    table.style = 'Light Shading'
    for index, col in enumerate(cols):
        table.cell(0, index).text = col
    for row in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = str(row)
    d.add_page_break()

    #                   Map of affected area
    add_heading('Map of affected area', level=2)

    return d
