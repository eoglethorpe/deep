from docx import Document, RT
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn

from entries.export_fields import *
from entries.models import InformationAttribute, InformationAttributeGroup, Sector

NO_SECTOR_LABEL = 'No Sector'

def _add_hyperlink(paragraph, url, text):
    """
    A work around function that places a hyperlink within a paragraph object.
    See: https://github.com/python-openxml/python-docx/issues/74#issuecomment-215678765

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: A Run object containing the hyperlink
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )
    hyperlink.set(qn('w:history'), '1')

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = OxmlElement('w:rPr')

    # Create a w:rStyle element, note this currently does not add the hyperlink style as its not in
    # the default template, I have left it here in case someone uses one that has the style in it
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')

    # Join all the xml elements together add add the required text to the w:r element
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return r    


def _sort(ents, order):
    """arrange entries based on hierarchy and place into ordered list
        we assume there are a maximum of 5 admin levels and all entries
        are given values for the admin levels"""
    #TODO: make work

    return ents

def _colorify_runs(runs, rgb_val):
    """take in a list of runs and make them the same color based on rgb"""
    for r in runs:
        f = r.font
        f.color.rgb = rgb_val


def _add_line(para):
    #a very convoluded way to add a horizontal line to the document
    #see: https://github.com/python-openxml/python-docx/issues/105

    def first_child_found_in(parent, tagnames):
        """
        Return the first child of parent with tag in *tagnames*, or None if
        not found.
        """
        for tagname in tagnames:
            child = parent.find(qn(tagname))
            if child is not None:
                return child
        return None

    def insert_element_before(parent, elm, successors):
        """
        Insert *elm* as child of *parent* before any existing child having
        tag name found in *successors*.
        """
        successor = first_child_found_in(parent, successors)
        if successor is not None:
            successor.addprevious(elm)
        else:
            parent.append(elm)
        return elm

    p = para._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    insert_element_before(pPr, pBdr, successors=(
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    ))
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    
def _xstr(v):
    """safely convert a value to string"""
    if v is None:
        return ''

    try:
        return str(v.encode('utf8').decode('UTF-8'))

    except:
        return str(v)

def _gen_bili(sortents):
    """generate bibliography
        return: list of oderededdicts with biliography entries:
        {
            Name of Source : name,
            Title of Lead : title,
            Publication Date : date(MM/DD/YYYY),
            Link: hyperlink (all in Title Case)
        }`
    """
    return [
        OrderedDict([
            ('Source' , get_source(e)),
            ('Title' , get_lead_nm_title_case(e)),
            ('Date' , get_lead_created_at_dt_num(e)),
            ('URL' , get_lead_url(e))
        ]) for e in sortents
    ]

def _add_bili(sortents, d):
    """insert bibliography into doc. Bib format:
    Name of Source. Title of Lead. Publication Date (MM/DD/YYYY). hyperlink (all in Title Case)
    UNHCR. Dont Have to Live Like a Refugee. 31/08/2016. https://mylink.com/
"""
    sep = d.add_paragraph()
    _add_line(sep)
    bib = d.add_paragraph()
    bib.add_run('Bibliography').bold = True

    for e in _gen_bili(sortents):
        d.add_paragraph()
        ent = d.add_paragraph()
        for i,v in enumerate(e.items()):
            if not v[1]:
                run = 'Missing ' + v[0]
            elif v[0] == 'URL':
                _add_hyperlink(ent, v[1], v[1] + ' ')
                run = ''
            else:
                run = _xstr(v[1])

            #to not add a period to last element
            if i < len(e.values())-1:
                ent.add_run(run + '. ')
            else:
                ent.add_run(run)

def _get_sev_shp(scr):
    BASE_DIR = 'static/img/doc_export/sev_'
    add = 0

    if scr == 'NOP':
        add = 1
    elif scr == 'MIN':
        add = 2
    elif scr == 'SOC':
        add = 3
    elif scr == 'SOM':
        add = 4
    elif scr == 'SEV':
        add = 5
    elif scr == 'CRI':
        add = 6

    return BASE_DIR + str(add) + '.png'

def _get_rel_shp(scr):
    BASE_DIR = 'static/img/doc_export/rel_'
    add = 0

    if scr == 'COM':
        add = 1
    elif scr == 'USU':
        add = 2
    elif scr == 'FAI':
        add = 3
    elif scr == 'NUS':
        add = 4
    elif scr == 'UNR':
        add = 5
    elif scr == 'CBJ':
        add = 6

    return BASE_DIR + str(add) + '.png'

c = 0
def _add_bod(doc, ent, att):
    """generate the text of an info att for a given entry"""
    p = doc.add_paragraph(style = 'gentext')

    #add excerpt
    p.add_run(get_ia_exc(att))

    #add source
    p.add_run(' (')
    if get_lead_url(ent):
        if not get_source(ent):
            _add_hyperlink(p, get_lead_url(ent), 'Reference')
        else:
            _add_hyperlink(p, get_lead_url(ent), get_source(ent))
    else:
        p.add_run('Manual Entry')

    #add date
    p.add_run(', ' + get_lead_created_at_dt_readable(ent) + ') ')

    #sev, rel
    r = p.add_run()
    #r.add_picture(_get_sev_shp(1), height=Inches(.5))
    r.add_picture(_get_sev_shp(att.severity), height=Inches(.2))
    r.add_text(' ')
    r.add_picture(_get_rel_shp(att.reliability), height=Inches(.2))

    doc.add_paragraph(style = 'gentext')

def _split_ent_by_sect(ents, sects):
    """take in entries and sects and split entries into provided sects (there can be repeats)
        return: {sector1: [ent1... entx], sector2: [ent1... entx]}
    """

    #create a list of [(IA, [parent sectors])]... set is used to avoid repeats
    divd = []
    for ent in ents:
        if len(ent.sectors.all()) == 0:
            divd.append((ent, [Sector(NO_SECTOR_LABEL)]))

        for v in ent.sectors.all():
            hold = []
            if v.parent == None:
                hold.append(v)
            else:
                hold.append(v.parent)

            divd.append((ent, list(set(hold))))

    ret = {}
    for att in divd:
        for sect in att[1]:
            if sect not in ret:
                ret[sect] = [att[0]]
            else:
                ret[sect].append(att[0])

    return ret

def _get_parents_sects():
    """get all parent sects"""
    #workaround to add a 'None' sector where no sectors are selected
    ret = [Sector(NO_SECTOR_LABEL)]

    #iterate through all sectors and pull their parents to get full list
    for sect in Sector.objects.all():
        #if a given subsector doesn't have a parent then we declare it a parent
        if not sect.parent:
            ret.append(sect)
        else:
            if sect.parent not in ret:
                ret.append(sect)

    return ret

def _add_sectors(d, ents, subcats):
    """take in all entries, divide them by sectors and print out"""
    parent_sects = _get_parents_sects()
    ents_by_sect = _split_ent_by_sect(ents, subcats)

    #go through each sector and print out IAs based on their categories
    for s in parent_sects:
        p = d.add_paragraph(style='mainhead')
        p.add_run(s.name)
        d.add_paragraph()
        d.add_paragraph(style = 'gentext')

        #add in info attribute group headings
        for cat in set(subcats):
            p = d.add_paragraph(style='secondhead')
            p.add_run(cat.name)
            d.add_paragraph(style = 'gentext')

            #go through the InformationAttributes that are in this particular group and write out IA name
            #...regardless of it is contained in data or not. if relevant IA is present, add in text
            for ia in cat.informationattribute_set.all():
                has_val = False
                p = d.add_paragraph(style='iatext')
                p.add_run(ia.name + ':')

                #now check our dict to see if relevant att is present with entries
                if s in ents_by_sect.keys():
                    for e in ents_by_sect[s]:
                        for iv in [att for att in e.attributedata_set.all() if att.attribute == ia]:
                            _add_bod(d, iv.entry, iv)

                if not has_val:
                    d.add_paragraph(style='gentext')

def _find_att(ents, att):
    """take in a list of entries and the desired InfoAttribute and return the attribute objects and entry"""
    ret = []
    for e in ents:
        for v in e.attributedata_set.all():
            if v.attribute == att:
                ret.append(v)

    return ret

def _add_gen_cat_second(d, ents, att):
    """add in secondary headers"""
    p = d.add_paragraph(style='secondhead')
    p.add_run(att.name)
    d.add_paragraph(style = 'gentext')

    has_val = False

    #iterate through all applicable IAs in a given IA group and get their values
    for v in _find_att(ents, att):
        _add_bod(d, v.entry, v)
        has_val = True

    if not has_val:
        d.add_paragraph(style = 'gentext')

def _add_gen_cat(d, ents, cat):
    p = d.add_paragraph(style='mainhead')
    p.add_run(cat.name)
    d.add_paragraph()
    d.add_paragraph(style = 'gentext')

    for att in [v for v in InformationAttribute.objects.all() if v.group == cat]:
        #if the given IA group is present in our data
        _add_gen_cat_second(d, ents, att)


def _add_meta(d):
    """just show sev and rel translations"""

def _gendoc(order, sortents):
    """generate the docx based on given order"""

    #read in a blank template that contains our desired fields
    d = Document('media/doc_export/template.docx')

    #meta info
    _add_meta(d)

    #the first part of export are universal categories that sector independant and the rest are sorted by sector
    single_cats = [InformationAttributeGroup('Context'), InformationAttributeGroup('Population data and characterstics'),
                   InformationAttributeGroup('Humanitarian access'),
                   InformationAttributeGroup('Information and communication')]
    all_cats = [v.group for v in InformationAttribute.objects.all()]

    for cat in single_cats:
        if cat not in all_cats:
            raise Exception('Given category is not present. Model may have changed.')
        else:
            _add_gen_cat(d, sortents, cat)

    #now we add data sorted by sector
    _add_sectors(d, sortents, [v for v in all_cats if v not in single_cats])

    #show bibliography
    _add_bili(sortents, d)

    return d

def export(order):
    return _gendoc(order, _sort(Entry.objects.all(), order))